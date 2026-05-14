import streamlit as st
import pandas as pd
import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime

# 1. Setup
st.set_page_config(page_title="St. Paul's Timetable Tool", layout="wide")
st.title("🏫 Teacher Timetable Assistant")

@st.cache_data
def load_data():
    try:
        df = pd.read_excel("master_timetable.xlsx", header=[2, 3])
        df = df.dropna(how='all', axis=0).dropna(how='all', axis=1)
        return df
    except Exception as e:
        st.error(f"Error loading Excel file: {e}")
        return None

# 2. High-Fidelity Word Document Generator
def generate_formal_docx(sender, receiver, target_class, s_day, s_per, r_day, r_per, reason):
    doc = Document()
    
    # Global Font Setup: Times New Roman 11pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Set Asian/ASCII font consistency
    r_pr = doc.styles['Normal']._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn('w:ascii'), 'Times New Roman')
    r_pr.get_or_add_rFonts().set(qn('w:hAnsi'), 'Times New Roman')

    # Set Paper Size to A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # --- Headings ---
    h1 = doc.add_paragraph("St. Paul’s School (Lam Tin)")
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.runs[0].bold = True
    h1.runs[0].font.size = Pt(14)
    h1.paragraph_format.space_after = Pt(2)

    h2 = doc.add_paragraph("Record of Exchange of Lessons")
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.runs[0].bold = True
    h2.runs[0].font.size = Pt(12)
    h2.paragraph_format.space_after = Pt(18)

    # --- Info Fields (Borderless Table with Bottom Borders and Vertical Alignment) ---
    def set_cell_bottom_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
            tcPr.append(tcBorders)
        bottom = parse_xml(r'<w:bottom %s w:val="single" w:sz="4" w:space="0" w:color="auto"/>' % nsdecls('w'))
        tcBorders.append(bottom)

    info_table = doc.add_table(rows=2, cols=2)
    info_table.autofit = False
    info_table.columns[0].width = Inches(1.6)
    info_table.columns[1].width = Inches(5.0)

    for row in info_table.rows:
        row.height = Cm(0.8)

    # Teacher Name Row
    cell_n_label = info_table.cell(0, 0)
    cell_n_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_n = cell_n_label.paragraphs[0]
    p_n.add_run("Name of Teacher:").bold = True

    cell_n_val = info_table.cell(0, 1)
    cell_n_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bottom_border(cell_n_val)
    cell_n_val.text = f" {sender}"

    # Reason Row
    cell_r_label = info_table.cell(1, 0)
    cell_r_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_r = cell_r_label.paragraphs[0]
    p_r.add_run("Reason for Exchange:").bold = True

    cell_r_val = info_table.cell(1, 1)
    cell_r_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bottom_border(cell_r_val)
    cell_r_val.text = f" {reason if reason else ''}"

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- Main Table (6 Rows: 2 Header + 4 Content) ---
    table = doc.add_table(rows=6, cols=14)
    table.style = 'Table Grid'
    
    # Set Height for all 4 Content Rows (Indices 2, 3, 4, 5) to 0.85cm
    for i in range(2, 6):
        table.rows[i].height = Cm(0.85)
    
    # Merge Header Row 1
    c1 = table.cell(0, 0).merge(table.cell(0, 6))
    c1.text = "Lessons to be substituted"
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.paragraphs[0].runs[0].bold = True
    
    c2 = table.cell(0, 7).merge(table.cell(0, 13))
    c2.text = "Lessons to be returned"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraphs[0].runs[0].bold = True

    # Header Row 2
    sub_headers = ["Date", "Day", "Class", "Period", "Subject on Timetable", "Subject Replacing the Original", "Name of Teacher Taking the Lesson"]
    full_headers = sub_headers + sub_headers
    for i, txt in enumerate(full_headers):
        cell = table.cell(1, i)
        cell.text = txt
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(8)

    def clean_p(v): return str(v).replace("P", "").strip()

    # Fill Data Row (Row index 2)
    table.cell(2, 1).text = str(s_day)
    table.cell(2, 2).text = str(target_class)
    table.cell(2, 3).text = clean_p(s_per)
    table.cell(2, 6).text = str(receiver)

    if r_day and r_day != "None":
        table.cell(2, 8).text = str(r_day)
        table.cell(2, 9).text = str(target_class)
        table.cell(2, 10).text = clean_p(r_per)
        table.cell(2, 13).text = str(sender)

    # --- Footer Section ---
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    today = datetime.now().strftime("%d / %m / %Y")
    
    ft = doc.add_table(rows=2, cols=2)
    ft.autofit = False
    ft.columns[0].width = Inches(3.5)
    ft.columns[1].width = Inches(3.5)

    ft.cell(0, 0).text = "Signature of teacher: ____________________"
    ft.cell(1, 0).text = f"Date: {today}"

    p_app = ft.cell(0, 1).paragraphs[0]
    p_app.text = "Approved by Principal: ____________________"
    p_app.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_dt = ft.cell(1, 1).paragraphs[0]
    p_dt.text = "Date: ____________________"
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 3. Main Logic
try:
    df = load_data()
    if df is not None:
        teacher_col = df.columns[0]
        teachers = sorted(df[teacher_col].dropna().unique().tolist())
        available_days = [d for d in df.columns.levels[0] if "Day" in str(d)]

        def is_free(val, disregards):
            val = str(val).strip().upper()
            if val in ["NAN", "", "NONE"]: return True
            for d in disregards:
                if "*" in d:
                    if val.startswith(d.replace("*", "").upper()): return True
                elif val == d.upper(): return True
            return False

        def teaches_class(t_name, target_class):
            row = df[df[teacher_col] == t_name].iloc[0]
            return target_class.upper() in [str(v).strip().upper() for v in row.values[1:]]

        tab1, tab2 = st.tabs(["🔍 Find free lesson 「Call會快」", "🔄 Swap Lesson 「調堂易」"])

        with tab1:
            st.header("Find Common Free Lessons")
            sel_t = st.multiselect("Select Teachers", teachers)
            sel_d = st.multiselect("Select Days", available_days, default=available_days)
            dis_in = st.text_input("Disregard (e.g. 6*, CLP)")
            dis_l = [x.strip() for x in dis_in.split(",") if x.strip()]

            if sel_t:
                results = []
                subset = df[df[teacher_col].isin(sel_t)]
                for day in sel_d:
                    for period in df[day].columns:
                        if all(is_free(row[(day, period)], dis_l) for _, row in subset.iterrows()):
                            results.append({"Day": day, "Period": f"P{period}"})
                if results:
                    st.table(pd.DataFrame(results).groupby('Day')['Period'].apply(lambda x: ", ".join(x)).reset_index())
                else: st.warning("No common free slots found.")

        with tab2:
            st.header("Swap Lesson Finder")
            c1, c2, c3 = st.columns(3)
            with c1: my_name = st.selectbox("Your Name", ["Select..."] + teachers)
            with c2: s_day = st.selectbox("Day of Lesson", available_days)
            with c3:
                p_list = list(df[s_day].columns) if s_day in available_days else []
                s_per = st.selectbox("Period", p_list)

            dis_sw = st.text_input("Disregard Classes (e.g. 6*)", key="sw_dis")
            dis_l_s = [x.strip() for x in dis_sw.split(",") if x.strip()]

            if my_name != "Select...":
                my_row = df[df[teacher_col] == my_name].iloc[0]
                target_class = str(my_row[(s_day, s_per)]).strip()

                # LOGIC UPDATED HERE: CLP disallowed similarly to M classes
                if is_free(target_class, dis_l_s):
                    st.error(f"You are FREE on {s_day} P{s_per}.")
                elif target_class.upper().endswith('M'):
                    st.error(f"Class {target_class} is a mixed (M) class. Cannot swap.")
                elif target_class.upper() == "CLP":
                    st.error(f"CLP lessons cannot be swapped.")
                else:
                    st.info(f"Finding swaps for **{target_class}** on {s_day} P{s_per}")
                    partners = []
                    for _, row in df.iterrows():
                        other = row[teacher_col]
                        if other == my_name: continue
                        if is_free(row[(s_day, s_per)], dis_l_s):
                            if teaches_class(other, target_class):
                                rets = []
                                for d in available_days:
                                    for p in df[d].columns:
                                        if str(row[(d, p)]).strip().upper() == target_class.upper():
                                            if is_free(my_row[(d, p)], dis_l_s):
                                                rets.append(f"{d} P{p}")
                                partners.append({"Colleague": other, "Returns": rets})

                    if partners:
                        view = pd.DataFrame([{"Colleague": p["Colleague"], "Returns": ", ".join(p["Returns"]) if p["Returns"] else "None"} for p in partners])
                        st.table(view)
                        st.divider()
                        st.subheader("📄 Generate Official Exchange Slip")
                        reason = st.text_input("Reason for Exchange")
                        ec1, ec2 = st.columns(2)
                        with ec1: sel_p = st.selectbox("Select Colleague", [p["Colleague"] for p in partners])
                        p_data = next(p for p in partners if p["Colleague"] == sel_p)
                        with ec2: sel_ret = st.selectbox("Select Return Lesson", p_data["Returns"] if p_data["Returns"] else ["None"])
                        
                        if st.button("Prepare Download"):
                            r_day, r_per = "None", "None"
                            if " " in sel_ret:
                                pts = sel_ret.split(" ")
                                r_day, r_per = f"{pts[0]} {pts[1]}", pts[2]
                            doc_bytes = generate_formal_docx(my_name, sel_p, target_class, s_day, s_per, r_day, r_per, reason)
                            st.download_button("⬇️ Download Lesson Exchange Slip", doc_bytes, f"Swap_{target_class}_{my_name}.docx")
                    else: st.warning("No partners found.")
except Exception as e:
    st.error(f"Error: {e}")
